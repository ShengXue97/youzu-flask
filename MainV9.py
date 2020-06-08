# V9: Current implementation - Working on Math Papers
import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches
from GibberishDetector import classify
import os
import sys
from pdf2image import convert_from_path
import io
import re
import pandas as pd
import platform
import math as m
import ast
import shutil
import time
import settings

def get_image(image_path):
    """Get a numpy array of an image so that one can access values[x][y]."""
    image = Image.open(image_path, 'r')
    image = image.convert('L')  # makes it greyscale
    width, height = image.size
    pixel_values = list(image.getdata())
    if image.mode == 'RGB':
        channels = 3
    elif image.mode == 'L':
        channels = 1
    else:
        print("Unknown mode: %s" % image.mode)
        return None
    pixel_values = np.array(pixel_values).reshape((width, height, channels))
    return pixel_values


def is_white_image(image_name):
    numpy_array = get_image(image_name + ".jpg")
    total_pixels = numpy_array.size
    num_of_black = 0
    num_of_white = 0

    for i in numpy_array:
        for j in i:
            if j[0] > 200:
                num_of_white = num_of_white + 1
            else:
                num_of_black = num_of_black + 1

    white_percentage = num_of_white / total_pixels
    # Save as inverted image if it is a negative image
    if white_percentage < 0.75:
        # load image# Load image
        im = Image.open(image_name + ".jpg")

        # Invert
        result = ImageOps.invert(im)

        # Save
        result.save(image_name + "_inverted.jpg")
        return False
    else:
        return True


def get_thresh_and_contours(img, filename):
    # Blurring
    imgBlur = cv2.GaussianBlur(img, (7, 7), 1)
    # cv2.imshow("blur",imgBlur)

    # convert to gray
    gray = cv2.cvtColor(imgBlur, cv2.COLOR_BGR2GRAY)
    # cv2.imshow("gray",gray)

    # threshold the grayscale image
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    # Draw contours
    result = img.copy()

    # use morphology erode to blur horizontally
    # kernel = np.ones((500,3), np.uint8)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (200, 3))  # 250,3
    morph = cv2.morphologyEx(thresh, cv2.MORPH_DILATE, kernel)

    if re.search(r'english', filename, re.I):
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (20, 30))  # previously 20,30 for eng papers
    else:
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 17))  # previously 20,30 for eng papers
    morph = cv2.morphologyEx(morph, cv2.MORPH_OPEN, kernel)

    resized = cv2.resize(morph, (700, 850))
    # cv2.imshow("morph",resized)

    # find contours
    cntrs = cv2.findContours(morph, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cntrs = cntrs[0] if len(cntrs) == 2 else cntrs[1]
    return thresh, cntrs, result, morph


def merge_contours(thresh, cntrs, x_tolerance, y_tolerance):
    # Erase small contours, and contours which small aspect ratio (close to a square)
    for c in cntrs:
        area = cv2.contourArea(c)

        # Fill very small contours with zero (erase small contours).
        if area < 100:
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue

        # https://stackoverflow.com/questions/52247821/find-width-and-height-of-rotatedrect
        rect = cv2.minAreaRect(c)
        (x, y), (w, h), angle = rect
        # aspect_ratio = max(w, h) / min(w, h)

        # Assume zebra line must be long and narrow (long part must be at lease 1.5 times the narrow part).
        '''
        if (aspect_ratio < 1.5):
            cv2.fillPoly(thresh, pts=[c], color=0)
            continue
            '''

    # Use "close" morphological operation to close the gaps between contours
    # https://stackoverflow.com/questions/18339988/implementing-imcloseim-se-in-opencv

    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE,
                              cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (x_tolerance, y_tolerance)));

    # Find contours in thresh_gray after closing the gaps
    cntrs, hier = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
    return thresh, cntrs


def draw_contours(result, img, cntrs, image_name):
    global diagram_count
    global texted
    global filename
    # Contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = []
    height, width, channels = img.shape

    if re.search(r'english', filename, re.I):
        result_1 = img.copy()
        # image preprocessing and find contours for blank line detection
        gray_1 = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh_1 = cv2.threshold(gray_1, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (15, 1))
        detected_lines = cv2.morphologyEx(thresh_1, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)

        cnts_line = cv2.findContours(detected_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        cnts_line = cnts_line[0] if len(cnts_line) == 2 else cnts_line[1]
        # draw contours for blank line
        for c in cnts_line:
            # cv2.drawContours(image, [c], -1, (0,0,255), 3)
            x1, y1, w1, h1 = cv2.boundingRect(c)
            if 0.05 < w1 / width < 0.2:
                cv2.rectangle(result, (x1, y1), (x1 + w1, y1 + h1), (255, 0, 0), 2)
                texted = cv2.putText(result_1, '(EMPTY)____', (x1, y1), cv2.FONT_HERSHEY_SIMPLEX,
                                     0.8, (0, 0, 0), 2, cv2.LINE_AA)
                texted = cv2.dilate(texted, np.ones((2, 2), np.uint8), iterations=1)
        for c in cntrs:
            area = cv2.contourArea(c) / 10000
            x, y, w, h = cv2.boundingRect(c)
            cv2.rectangle(result, (x - 10, y - 10), (x + w + 5, y + h + 5), (0, 0, 255), 2)
            if "texted" in locals():
                grayish = cv2.cvtColor(texted, cv2.COLOR_BGR2GRAY)
                #image = cv2.imread(img, 0)
            else:
                grayish = cv2.cvtColor(result_1, cv2.COLOR_BGR2GRAY)
                # image = cv2.imread(img, 0)
            thresh = 255 - cv2.threshold(grayish, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1][1]
            ###cropping out image and convert to text
            ROI = thresh[y:y + h, x:x + w]
            text = pytesseract.image_to_string(ROI, lang='eng', config='--psm 6')
            text = re.sub(r"\(EMPTY[\)]*|\(FMPTY[\)]*", "_________", text)
            pseudo_text = text

            # Only add the image if it is large enough,
            # and the entire image has illegal text symbols(which is likely to be a diagram)
            # if w/width > 0.0302 and h/height > 0.0213 and y/height < 0.95:
            if w / width > 0.05 and y / height < 0.95:  # and (x/width < 0.4 or x/width > 0.5)
                # hough line detector included too for eng
                new_image = img[y:y + h, x:x + w]
                dst = cv2.Canny(new_image, 50, 200, None, 3)
                linesp = cv2.HoughLinesP(dst, 1, np.pi / 180, 50, None, 50, 2)
                # if is_gibberish(text) and w/h < 5:
                if is_gibberish(text) or 0.35 < (w * h) / (width * height) < 0.97 or linesp is not None:
                    if h / height > 0.1 and w / h < 10:
                        # Likely to be an image
                        new_image = img[y:y + h, x:x + w]
                        cv2.imwrite("TempImages/" + str(diagram_count) + ".jpg", new_image)
                        document_data_list.append(
                            ("TempImages/" + str(diagram_count) + ".jpg", "image", y, pseudo_text))
                        diagram_count = diagram_count + 1
                    else:
                        # Likely to be text, just small regions like "Go on to the next page"
                        document_data_list.append((text, "text", y, pseudo_text))
                else:
                    # Likely to be a text
                    document_data_list.append((text, "text", y, pseudo_text))
    # for non english papers, no blank line detection required
    else:
        for c in cntrs:
            area = cv2.contourArea(c) / 10000
            x, y, w, h = cv2.boundingRect(c)
            cv2.rectangle(result, (x - 10, y - 10), (x + w + 5, y + h + 5), (0, 0, 255), 2)
            if platform.system() == "Windows":
                pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"

            # Read as binary image
            image = cv2.imread(image_name + ".jpg", 0)

            thresh = 255 - cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

            ROI = thresh[y:y + h, x:x + w]
            text = pytesseract.image_to_string(ROI, lang='eng', config='--psm 6')
            pseudo_text = text

            # Only add the image if it is large enough,
            # and the entire image has illegal text symbols(which is likely to be a diagram)
            # if w/width > 0.0302 and h/height > 0.0213 and y/height < 0.95:
            if w / width > 0.1 and h / height > 0.001 and y / height < 0.95 and w / h < 25:  # and (x/width < 0.4 or x/width > 0.5)
                # if is_gibberish(text) and w/h < 5:
                ############## trying out hough transform as a filter!!! ###############
                new_image = img[y:y + h, x:x + w]
                dst = cv2.Canny(new_image, 50, 200, None, 3)
                linesp = cv2.HoughLinesP(dst, 1, np.pi / 180, 50, None, 50, 2)
                if linesp is not None or is_gibberish(text):
                    cv2.imwrite("TempImages/" + str(diagram_count) + ".jpg", new_image)
                    document_data_list.append(("TempImages/" + str(diagram_count) + ".jpg", "image", y, pseudo_text))
                    diagram_count = diagram_count + 1
                else:
                    # large chunk that resembles diagram, but really is text
                    document_data_list.append((text, "text", y, pseudo_text))
    return document_data_list


def is_gibberish(text):
    global filename
    is_definitely_not_gibberish = False
    split = text.split("\n")

    total_value = 0
    if len(split) == 0:
        return False

    # Every line in the contour box
    for s in split:
        if re.search(r'english', filename, re.I):
            if "(1)" in s or "(2)" in s or "(3)" in s or "(4)" in s or "(a)" in s or "(b)" in s or "(c)" in s or "(d)" in s:
                # if there is a answer number, it is definitely not gibberish. i.e we want it as a text, not image
                is_definitely_not_gibberish = True
        elif re.search(r'math', filename, re.I):
            # some qn headers for non-english papers are auto recognized as gibberish, so we want to correct that
            search_sentence = re.search(r'[?0-9]+[.,]+', s, re.I)
            numberans_search = re.search(r'[\[\(\{][1-4a-b]+[]\)\}]|Ans]', s, re.I)
            if search_sentence or numberans_search:
                is_definitely_not_gibberish = True
        else:
            # definitely a science paper, where options (1)-(4) should not always be treated as not gibberish
            search_sentence = re.search(r'[?0-9]+[.,]+', s, re.I)
            if search_sentence or "(a)" in s or "(b)" in s or "(c)" in s or "(d)" in s:
                is_definitely_not_gibberish = True

    gibberish_likelihood_percentage = classify(s)
    total_value = total_value + gibberish_likelihood_percentage

    if is_definitely_not_gibberish:
        return False

    average_percentage = total_value / len(split)
    if average_percentage > 45:
        # likely to be gibberish
        return True
    else:
        return False
    

def write_data_to_document(document_data_list, document, qn_coord):
    # qn_coord is a tuple consisting of nested (pg_number, y) tuples of each contour
    global qn_num
    global pg_num
    global global_df
    global current_ans_list
    global found_ans_options
    global file_attribute_list
    global image_count

    # Sort data of text and images according to their y values, and add them to a word document
    document_data_list.sort(key=lambda tup: tup[2])

    for i in range(len(document_data_list)):
        data = document_data_list[i]
        item = data[0]
        typeof = data[1]
        y_coord = data[2]
        pseudo_text = data[3]
        # Eg. ("1.jpg", "image", "45", "@#!$@!$@!$")
        # STEP 1: Add qn number if current contour exceeds the page number or y_coord of the current question
        # (pg_num, y_coord, qn_num, qn_section)
        if qn_num < len(qn_coord) - 1:
            next_qn_tuple = qn_coord[qn_num + 1]
            # print((pg_num, y_coord, next_qn_tuple[0], next_qn_tuple[1]))
            if pg_num > next_qn_tuple[0] or (pg_num == next_qn_tuple[0] and y_coord > (next_qn_tuple[1] - 5)):
                qn_num = qn_num + 1
                current_ans_list = []
                found_ans_options = False

        # STEP 2: Find ans sections
        regex = re.compile('[\[\(\|\{][0-9][\]\)\}\|]')

        matches = regex.finditer(pseudo_text)
        match_list = []
        for match in matches:
            match_list.append(match)

        first_ans_pos = -1
        for i in range(len(match_list)):
            match = match_list[i]
            lineno = pseudo_text.count('\n', 0, match.start())

            if len(current_ans_list) <= 3:
                startpos = match.regs[0][0]
                endpos = match.regs[0][1]
                if len(current_ans_list) == 0:
                    first_ans_pos = startpos

                if i == len(match_list) - 1:
                    # Last match
                    substr = pseudo_text[startpos:]
                    current_ans_list.append(substr)
                else:
                    # Still have matches after
                    nextstartpos = match_list[i + 1].regs[0][0]
                    substr = pseudo_text[startpos:nextstartpos]
                    current_ans_list.append(substr)
            else:
                break

        ans_a = "-" if len(current_ans_list) <= 0 else current_ans_list[0]
        ans_b = "-" if len(current_ans_list) <= 1 else current_ans_list[1]
        ans_c = "-" if len(current_ans_list) <= 2 else current_ans_list[2]
        ans_d = "-" if len(current_ans_list) <= 3 else current_ans_list[3]

        # file_attribute_list -> [paper_level, paper_subject, paper_year, paper_school, paper_exam_type]
        paper_level = file_attribute_list[0].upper()
        paper_subject = file_attribute_list[1].upper()
        paper_year = file_attribute_list[2]
        paper_school = file_attribute_list[3]
        paper_exam_type = file_attribute_list[4].upper()

        # STEP 3: Add question to dataframe
        if typeof == "text":
            document.add_paragraph(item)
            illegal_qn_strings = ["chij", "mark", "instructions", "go on to the next page", "blank page", "question"]
            # Do not accept text as question if it contains any of these strings
            contains_illegal_qn_string = any(ele in item.lower() for ele in illegal_qn_strings)
            mcq_identifiers = ["(1)", "(2)", "(3)", "(4)"]
            contains_mcq_identifier = any(ele in item.lower() for ele in mcq_identifiers)
            if not contains_illegal_qn_string and item != "":
                # print((pg_num, y_coord, qn_num, item))
                # ['Level', 'Question', 'isMCQ', 'A', 'B', 'C', 'D', 'Subject', 'Year', 'School', 'Exam', 'Number', 'Image', 'Image File' ]
                if qn_num not in global_df.index:
                    if contains_mcq_identifier or ans_a != "" or ans_b != "" or ans_c != "" or ans_d != "":
                        global_df.loc[qn_num] = [paper_level, pg_num, item, "MCQ", ans_a, ans_b, ans_c, ans_d,
                                                 paper_subject,
                                                 paper_year, paper_school, paper_exam_type, qn_num, "No", "-"]
                    else:
                        global_df.loc[qn_num] = [paper_level, pg_num, item, "-", ans_a, ans_b,
                                                 ans_c, ans_d, paper_subject,
                                                 paper_year, paper_school, paper_exam_type, qn_num, "No", "-"]

                else:
                    if len(current_ans_list) != 0 and not found_ans_options:
                        # If an answer option was found in this contour (So, remove the answer options from the question)
                        item = item[:first_ans_pos]
                        found_ans_options = True
                    elif len(current_ans_list) != 0 and found_ans_options:
                        # If an answer option was found in a contour before(Means no more text in the actual question)
                        item = ""

                    global_df.loc[qn_num, 'Question'] = global_df.loc[qn_num, 'Question'] + item
                    global_df.loc[qn_num, 'A'] = ans_a
                    global_df.loc[qn_num, 'B'] = ans_b
                    global_df.loc[qn_num, 'C'] = ans_c
                    global_df.loc[qn_num, 'D'] = ans_d

        elif typeof == "image":
            if qn_num not in global_df.index:
                global_df.loc[qn_num] = [paper_level, pg_num, "-", "-", ans_a, ans_b, ans_c, ans_d, paper_subject, paper_year,
                                         paper_school, paper_exam_type, qn_num, "Yes", item]
            else:
                global_df.loc[qn_num, 'A'] = ans_a
                global_df.loc[qn_num, 'B'] = ans_b
                global_df.loc[qn_num, 'C'] = ans_c
                global_df.loc[qn_num, 'D'] = ans_d
                global_df.loc[qn_num, 'Image'] = "Yes"
                global_df.loc[qn_num, 'Image File'] = global_df.loc[qn_num, 'Image File'] + ";" + item

            document.add_picture(item, width=Inches(5))
            image_count += 1


def generate_document(imagefilename, documentdir, qn_coord, requestID):
    global pg_num
    global total_pages
    global filename

    print("Step 2 (Output Generation): PG " + str(pg_num) + "/" + str(total_pages))
    entry = {'stage': 2, 'page' : pg_num, 'total' : total_pages, 'output' : []}
    settings.db[requestID] = entry

    time.sleep(0)

    image_name = imagefilename.replace(".jpg", "")
    document = Document()

    ###### Step 1: Convert to positive image if image is negative
    if not is_white_image(image_name):
        image_name = image_name + "_inverted"

    ###### Step 2: Get the initial thresh and contours
    img = cv2.imread(image_name + ".jpg")
    height, width, channels = img.shape
    # if math paper, remove the vertical lines on the right side of the page
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    remove_vertical = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
    cnts = cv2.findContours(remove_vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = cnts[0] if len(cnts) == 2 else cnts[1]
    if re.search(r'math', filename, re.I):
        for c in cnts:
            x, y, w, h = cv2.boundingRect(c)
            if x / width > 0.7 or x / width < 0.1:
                cv2.drawContours(img, [c], -1, (255, 255, 255), 20)
    else:
        # eliminate unwanted horizontal lines at paper margins due to scanning
        for c in cnts:
            x, y, w, h = cv2.boundingRect(c)
            if x / width > 0.9 or x / width < 0.1:
                cv2.drawContours(img, [c], -1, (255, 255, 255), 20)
    thresh, cntrs, result, morph = get_thresh_and_contours(img, imagefilename)

    ###### Step 3: Merge contours that are close together
    # Modify the x and y tolerance to change how far it must be before it will merge!
    if re.search(r'english', filename, re.I):
        x_tolerance = m.floor(0.18138 * width)  # previously 300px 0.3
        y_tolerance = m.floor(0.014964 * height)  # previously 35px 0.02
    else:
        x_tolerance = m.floor(0.18 * width)  # previously 300px
        y_tolerance = m.floor(0.009 * height)  # previously 0.014964 #SX: 0.015

    thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)

    ###### Step 4: Draw the contours on the image
    # ordered_value_tuples contains ordered tuples of (text, y_coord)
    # document_data_list contains list of tuples of (data, type, y_coord)
    # data contains actual string if it is a text, and the image path in TempImages if it contains an image.
    # type is "text" or "image"
    # y_coord contains y coordinates of the text or image
    document_data_list = draw_contours(result, img, cntrs, image_name)
    # cv2.imwrite("contour_img/" + str(file_count) + ".jpg", result)

    ###### Step 5: Write and Save to a new Microsoft Word Document
    write_data_to_document(document_data_list, document, qn_coord)
    # Remove /images from image_name. example image_name is images/P6_English_2019_CA1_CHIJ/pg_1_P6_English_2019_CA1_CHIJ.jpg
    image_name = image_name.split('/', 1)[1]
    # Test paper name found in /images, example parentdir is P6_English_2019_CA1_CHIJ
    parentdir = image_name.split('/', 1)[0]

    # put this under documentdir
    if not os.path.exists(documentdir + "/" + parentdir):
        os.makedirs(documentdir + "/" + parentdir)

    document.save(documentdir + "/" + image_name + ".docx")

    # cv2.imshow("THRESH", thresh)
    # cv2.imshow("MORPH", morph)

    ####### Step 6: Display results
    ims = cv2.resize(result, (700, 850))
    cv2.imwrite("TempContours/" + str(pg_num) + ".jpg", ims)
    pg_num = pg_num + 1
    # cv2.imshow("RESULT", ims)
    # cv2.waitKey(0)
    # cv2.destroyAllWindows()


# Copies all files from src directory to dest directory
def copytree(src, dst, symlinks=False, ignore=None):
    if not os.path.exists(dst):
        os.makedirs(dst)
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            copytree(s, d, symlinks, ignore)
        else:
            if not os.path.exists(d) or os.stat(s).st_mtime - os.stat(d).st_mtime > 1:
                shutil.copy2(s, d)


# Map the page number and y coordinates of each question
def find_qn_coords(filenames_list, requestID):
    global total_pages
    qn_coord = []
    qn_coord.append((0, 0))
    pg_number = 1
    qn_num = 1
    diagram_count = 1

    for filename in filenames_list:
        print("Step 1 (Preprocessing): PG " + str(pg_number) + "/" + str(total_pages))
        entry = {'stage': 1, 'page' : pg_num, 'total' : total_pages, 'output' : []}
        settings.db[requestID] = entry

        time.sleep(0)

        # if pg_num < 24:
        #     continue
        image_name = filename.replace(".jpg", "")
        # image_name = filename.replace(".jpg", "")
        ###### Step 1: Convert to positive image if image is negative
        if not is_white_image(image_name):
            image_name = image_name + "_inverted"
        ###### Step 2A: Read the image and check for special sections
        img = cv2.imread(image_name + ".jpg")
        target_word = []
        sorted_cntr_tuples = []
        # usually only need to catch questions past Q20, due to weird question number e.g. (29)
        for a in range(21, 101):
            target_word.append('(' + str(a) + ')')
        # same case for strangely placed qns, but they have a "." in front instead of parenthesis
        for b in range(0, 101):
            target_word.append(str(b) + '.')

        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        for ele in target_word:
            word_occurences = [i for i, word in enumerate(data["text"]) if word.lower() == ele]
            for occ in word_occurences:
                w = data["width"][occ]
                h = data["height"][occ]
                x = data["left"][occ]
                y = data["top"][occ]
                sorted_cntr_tuples.append(("", pg_number, y, w, h, x))


        ###### Step 2B: Get the initial thresh and contours
        #img = cv2.imread(image_name + ".jpg")
        height, width, channels = img.shape
        thresh, cntrs, result, morph = get_thresh_and_contours(img, filename)

        ###### Step 3: Merge contours that are close together
        # Modify the x and y tolerance to change how far it must be before it will merge!
        x_tolerance = m.floor(0.02138 * width)  # previously 0.02138
        y_tolerance = m.floor(0.024964 * height)  # previously 0.024964
        thresh, cntrs = merge_contours(thresh, cntrs, x_tolerance, y_tolerance)

        for c in cntrs:
            area = cv2.contourArea(c) / 10000
            x, y, w, h = cv2.boundingRect(c)
            cv2.rectangle(result, (x - 10, y - 10), (x + w + 5, y + h + 5), (0, 0, 255), 2)
            if area < 0.1 and area > 0.01 and y / height < 0.855 and x / width < 0.24 and w / h < 2 and w / h > 0.5:
                new_image = img[y:y + h, x:x + w]
                text = pytesseract.image_to_string(new_image, lang='eng', config='--psm 6')
                if text != "":
                    # Check if pseudo_text contains numbers contained in any brackets
                    # matches = re.search(r'[\[\(\|\{][0-9a-z][\]\)\}\|]', text, re.I)
                    illegal_qn_strings = ["(", ")", "[", "]", "{", "}", "|", "NO"]
                    # Do not accept text as question if it contains any of these strings
                    contains_illegal_qn_string = any(ele in text.lower() for ele in illegal_qn_strings)
                    # Do not repeat tuples in second round of qn finding
                    probably_same_contour = False
                    for onetuple in sorted_cntr_tuples:
                        if abs(y - onetuple[2]) < 15 and pg_number == onetuple[1]:
                            probably_same_contour = True
                    if not contains_illegal_qn_string and not probably_same_contour:
                        sorted_cntr_tuples.append((c, pg_number, y, w, h, x / width))

        sorted_cntr_tuples.sort(key=lambda tup: tup[2])
        # Comment this out to visualise the processing of small contours under TempImages/small_[NAME].jpg
        small_cntrs = []
        for c, pg_number, y, w, h, xw in sorted_cntr_tuples:
            if c is not "":
                x, y, w, h = cv2.boundingRect(c)
                new_image = img[y:y + h, x:x + w]
            else:
                new_image = img[y:y + h, xw:xw + w]
            cv2.imwrite("TempImages/small_" + str(diagram_count) + ".jpg", new_image)
            # Read as binary image - only if we want to test what small image was read as (DEBUGGING)
            small_image = cv2.imread("TempImages/small_" + str(diagram_count) + ".jpg", 0)
            small_cntrs.append((pg_number, y))
            diagram_count = diagram_count + 1

        for c, pg_number, y, w, h, xw in sorted_cntr_tuples:
            qn_coord.append((pg_number, y))
            # print((pg_num, y))
        # page number will increment with every /small_ image appended to TempContours/
        cv2.imwrite("TempContours/" + str(pg_number) + ".jpg", result)
        pg_number = pg_number + 1

    return qn_coord


def find_paper_attributes(paper_name):
    global filename
    paper_name = paper_name.lower()
    paper_subject = ""
    paper_level = ""
    paper_exam_type = ""
    paper_year = ""
    paper_school = ""

    if re.search(r'english', paper_name, re.I):
        paper_subject = "english"
    elif re.search(r'math', paper_name, re.I):
        paper_subject = "math"
    elif re.search(r'science', paper_name, re.I):
        paper_subject = "science"

    if re.search(r'p[0-9]', paper_name, re.I):
        match = re.search(r'p[0-9]', paper_name, re.I)
        startpos = match.regs[0][0]
        endpos = match.regs[0][1]
        paper_level = paper_name[startpos:endpos]

    if re.search(r'ca1', paper_name, re.I):
        paper_exam_type = "ca1"
    elif re.search(r'ca2', paper_name, re.I):
        paper_exam_type = "ca2"
    elif re.search(r'sa1', paper_name, re.I):
        paper_exam_type = "sa1"
    elif re.search(r'sa2', paper_name, re.I):
        paper_exam_type = "sa2"

    if re.search(r'[0-9][0-9][0-9][0-9]', paper_name, re.I):
        match = re.search(r'[0-9][0-9][0-9][0-9]', paper_name, re.I)
        startpos = match.regs[0][0]
        endpos = match.regs[0][1]
        paper_year = paper_name[startpos:endpos]

    paper_name_split = paper_name.split("_")
    if len(paper_name_split) <= 1:
        paper_name_split = paper_name.split("-")

    illegal_paper_name_strings = [paper_subject, paper_level, paper_exam_type, paper_year]
    for i in illegal_paper_name_strings:
        if i == "":
            illegal_paper_name_strings.remove(i)

    for part in paper_name_split:
        contains_illegal_paper_name_string = any(ele in part.lower() for ele in illegal_paper_name_strings)
        if not contains_illegal_paper_name_string:
            paper_school = paper_school + part + "_"

    paper_school = paper_school[:-1]
    return paper_level, paper_subject, paper_year, paper_school, paper_exam_type


def acc_matrix(image_count, verifier, pdfname):
    if any(verifier["Paper Name"].values == pdfname):
        num_qns = int(verifier.loc[pdfname, "Questions"])
        num_images = int(verifier.loc[pdfname, "Images"])
        qn_acc = (qn_num / num_qns) * 100
        img_acc = (image_count / num_images) * 100
        return qn_acc, img_acc
    else:
        pass


def main(pdfname, requestID):
    global total_pages
    global global_df
    global file_attribute_list
    print(pdfname)

    if not os.path.exists("TempImages"):
        os.makedirs("TempImages")

    if not os.path.exists("TempContours"):
        os.makedirs("TempContours")

    paper_name = pdfname.replace(".pdf", "")
    pdf_path = "ReactPDF/" + paper_name + ".pdf"
    pages = convert_from_path(pdf_path)
    pg_cntr = 1
    filenames_list = []
    file_attribute_list = find_paper_attributes(paper_name)

    sub_dir = str("images/" + pdf_path.split('/')[-1].replace('.pdf', '') + "/")
    if not os.path.exists(sub_dir):
        os.makedirs(sub_dir)


    for page in pages:
        filename = "pg_" + str(pg_cntr) + '_' + pdf_path.split('/')[-1].replace('.pdf', '.jpg')
        page.save(sub_dir + filename)
        pg_cntr = pg_cntr + 1
        filenames_list.append(sub_dir + filename)

    total_pages = len(filenames_list)
    qn_coord = find_qn_coords(filenames_list, requestID)
    # qn_coord = ast.literal_eval("[(0, 0, 0, 0), (2, 1365, 1, ''), (2, 1703, 2, ''), (3, 1131, 3, ''), (3, 1819, 4, ''), (4, 966, 5, ''), (4, 1779, 6, ''), (5, 2056, 7, ''), (6, 744, 8, ''), (6, 1934, 9, ''), (7, 757, 10, ''), (7, 1924, 11, ''), (8, 905, 12, ''), (8, 1710, 13, ''), (9, 1077, 14, ''), (9, 1914, 15, ''), (10, 1256, 16, ''), (11, 1630, 17, ''), (12, 1385, 18, ''), (13, 1432, 19, ''), (14, 1401, 20, ''), (15, 1292, 21, ''), (16, 1047, 22, ''), (17, 1295, 23, ''), (18, 1169, 24, ''), (19, 1005, 25, ''), (19, 1527, 26, ''), (20, 1535, 27, ''), (21, 1186, 28, ''), (21, 1968, 29, ''), (24, 1630, 30, 2), (26, 1591, 31, 2), (27, 1584, 32, 2), (29, 268, 33, 3), (30, 1935, 34, 2), (31, 1858, 35, 3), (32, 1035, 36, 3), (33, 1711, 37, 1), (34, 1553, 38, 1), (35, 1395, 39, 1), (36, 1739, 40, 3)]")

    for filename in filenames_list:
        generate_document(filename, "OutputDocuments", qn_coord, requestID)

    global_df.to_csv(requestID + "_output.csv")

    # Copies all the output to a new folder under Output/PDF NAME
    dirpath = os.getcwd()
    copytree(dirpath + "/TempContours", dirpath + "/Output/" + paper_name + "/TempContours")
    copytree(dirpath + "/TempImages", dirpath + "/Output/" + paper_name + "/TempImages")
    copytree(dirpath + "/images/" + paper_name, dirpath + "/Output/" + paper_name + "/images")
    #shutil.copyfile(dirpath + requestID + "_output.csv", dirpath + "/Output/" + paper_name + requestID + "_output.csv")

    global_df = pd.DataFrame(
        columns=['Level', 'Page', 'Question', 'Comment', 'A', 'B', 'C', 'D', 'Subject', 'Year', 'School', 'Exam', 'Number',
                 'Image', 'Image File'])
    shutil.rmtree(dirpath + "/TempContours")
    shutil.rmtree(dirpath + "/TempImages")
    shutil.rmtree(dirpath + "/images/")

    # Create an empty list 
    row_json = []
    
    # Iterate over each row 
    for index, rows in global_df.iterrows(): 
        # Create list for the current row 
        my_list =[rows["Level"], rows["Question"], rows["isMCQ"], rows["A"], rows["B"], rows["C"], rows["D"], rows["Subject"], rows["Year"], rows["School"], rows["Exam"], rows["Number"], rows["Image"], rows["Image File"]] 
        # append the list to the final list 
        row_json.append(my_list)

    entry = {'stage': 3, 'page' : 0, 'total' : 0, 'output' : row_json}
    settings.db[requestID] = entry


qn_num = 1
pg_num = 1
diagram_count = 1
filename = ""
total_pages = -1
image_count = 0
current_section = ""
current_ans_list = []
file_attribute_list = []
found_ans_options = False
global_df = pd.DataFrame(
    columns=['Level', 'Page', 'Question', 'Comment', 'A', 'B', 'C', 'D', 'Subject', 'Year', 'School', 'Exam', 'Number', 'Image',
             'Image File'])

# for curFilename in os.listdir("Sample Resources"):
#     if curFilename.endswith("P6_2019_English_SA1_Catholic_High.pdf"):
#         filename = curFilename
#         main(curFilename)
#         qn_num = 1
#         pg_num = 1
#         diagram_count = 1
#         total_pages = -1
#         image_count = 0
#         current_section = ""
#         current_ans_list = []
#         found_ans_options = False